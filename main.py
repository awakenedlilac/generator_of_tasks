"""Generator of tasks"""
import re
import random
from docx import Document
from docx.shared import Pt
import pymorphy2
import nltk
from nltk.tokenize import word_tokenize
nltk.download('punkt')

with open('/texts.txt',
          'r', encoding='utf-8') as f:
    texts = f.read()


class TextProcessor:

    """processes the texts"""

    def __init__(self):
        self.original_text = []

    def tokenized_text(self, all_texts):

        """splits the texts by # and removes the punctuation"""

        if not isinstance(all_texts, str):
            return None
        text = re.split(r'[#\n]', all_texts)
        for elem in text:
            if elem != '':
                self.original_text.append(re.split(r'[.?!]', elem))
        return self.original_text

    def return_original(self, _original_text):

        """return original texts"""

        self.tokenized_text(texts)
        return self.original_text

class Generator:

    """in charge of the tasks"""

    def __init__(self, text: TextProcessor):
        self.text = text
        self.original_texts = self.getting_texts()

    def getting_texts(self):

        """takes 6 texts"""

        original_texts = random.sample(self.text.return_original(texts), 6)
        return original_texts

    def task_1(self, text):

        """shuffles the words in the sentence"""

        words = []
        for sentence in text:
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                word = word_tokenize(sentence)
                random.shuffle(word)
                words.append(word)
        return '. '.join(' '.join(word) for word in words)

    def task_2(self, text):

        """changes the verbs on infinitives"""

        morph = pymorphy2.MorphAnalyzer()
        sentences = []
        for sentence in text:
            words = []
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                for word in word_tokenize(sentence.lower()):
                    elem = morph.parse(word)[0]
                    if elem.tag.POS == 'VERB':
                        word = morph.parse(word)[0].normal_form
                    words.append(word)
                sentences.append(' '.join(words).capitalize())
        return '. '.join(sentences)

    def task_3(self, text):

        """makes halves of the sentences"""

        words = []
        for sentence in text:
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                words.append(word_tokenize(sentence))
        parts_1 = []
        parts_2 = []
        for sent in words:
            parts_1.append(' '.join(sent[:len(sent) // 2]))
            parts_2.append(' '.join(sent[len(sent) // 2:]))
        random.shuffle(parts_2)
        exercise_3 = []
        for i in zip(parts_1, parts_2):
            exercise_3.append(list(i))
        return exercise_3

    def task_4(self, text):

        """removes random words from the sentence"""

        words = []
        for sentence in text:
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                words.append(word_tokenize(sentence))
        answers = []
        final = []
        counter = 1
        for sentence in words:
            index = random.randint(0, len(sentence) - 1)
            if sentence[index].isalpha():
                answers.append(sentence[index].lower())
                sentence[index] = f'({counter})'
                counter += 1
                final.append(' '.join(sentence))
        random.shuffle(answers)
        exercise_4 = ['. '.join(final), answers]
        return exercise_4

    def task_5(self, text):

        """removes punctuation"""

        words = []
        for sentence in text:
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                words.append(word_tokenize(sentence))
        for sentence in words:
            for word in sentence:
                if word in ',?!:;':
                    sentence.remove(word)
        return '. '.join(' '.join(word) for word in words)

    def task_6(self, text):

        """removes random words from the sentence where we need to put an accent"""

        words = []
        for sentence in text:
            if not isinstance(sentence, str) or not sentence:
                text.remove(sentence)
            else:
                words.append(word_tokenize(sentence))
        answers = []
        for sentence in words:
            index = random.randint(0, len(sentence) - 1)
            if len(sentence[index]) > 3:
                answers.append(sentence[index])
        return ', '.join(answers)


class Storage:

    """storages the tasks and original texts to docx files"""

    def __init__(self, saved_task):
        self.saved_task = saved_task
        self.original_texts = self.saved_task.getting_texts()
        self.doc_orig = Document()
        self.doc = Document()
        self.save_original_texts()
        self.save_task_1()
        self.save_task_2()
        self.save_task_3()
        self.save_task_4()
        self.save_task_5()
        self.save_task_6()
        self.storage_all(self.doc_orig, self.doc)

    def save_original_texts(self):

        """saves original texts"""

        style = self.doc_orig.styles['Normal']
        style.font.name = 'Times New Roman'
        self.doc_orig.add_paragraph('\n\n'.join('. '.join(x)
                                               for x in self.original_texts))

    def save_task_1(self):

        """saves the first task"""

        para = self.doc.add_paragraph().add_run('Первое задание: '
                                                'поставьте слова в правильном порядке.')
        para.bold = True
        para.font.size = Pt(14)
        self.doc.add_paragraph(self.saved_task.task_1(self.original_texts[0]))

    def save_task_2(self):

        """saves the second task"""

        para = self.doc.add_paragraph().add_run('Второе задание: '
                                                'поставьте глаголы в нужную по контексту форму.')
        para.bold = True
        para.font.size = Pt(14)
        self.doc.add_paragraph(self.saved_task.task_2(self.original_texts[1]))

    def save_task_3(self):

        """saves the third task"""

        para = self.doc.add_paragraph().add_run('Третье задание: соедините части предложения.')
        para.bold = True
        para.font.size = Pt(14)
        table_3 = self.doc.add_table(rows=1, cols=2)
        table_3.style = 'Table Grid'
        row_3 = table_3.rows[0].cells
        row_3[0].text = 'Начало'
        row_3[1].text = 'Конец'
        for parts in self.saved_task.task_3(self.original_texts[2]):
            row_3 = table_3.add_row().cells
            row_3[0].text = parts[0]
            row_3[1].text = parts[1]

    def save_task_4(self):

        """saves the fourth task"""

        para = self.doc.add_paragraph().add_run('Четвертое задание: вставьте слова.')
        para.bold = True
        para.font.size = Pt(14)
        table_4 = self.doc.add_table(rows=1, cols=2)
        table_4.style = 'Table Grid'
        row_4 = table_4.rows[0].cells
        row_4[0].text = 'Слово'
        row_4[1].text = 'Номер'
        for elem in self.saved_task.task_4(self.original_texts[3]):
            if isinstance(elem, str):
                self.doc.add_paragraph('\n' + elem)
            else:
                for word in elem:
                    row_4 = table_4.add_row().cells
                    row_4[0].text = word
                    row_4[1].text = ''

    def save_task_5(self):

        """saves the first task"""

        para = self.doc.add_paragraph().add_run('Пятое задание: поставьте '
                                                'знаки препинания.')
        para.bold = True
        para.font.size = Pt(14)
        self.doc.add_paragraph(self.saved_task.task_5(self.original_texts[0]))

    def save_task_6(self):

        """saves the first task"""

        para = self.doc.add_paragraph().add_run('Шестое задание: поставьте ударение в словах.')
        para.bold = True
        para.font.size = Pt(14)
        self.doc.add_paragraph(self.saved_task.task_6(self.original_texts[0]))

    def storage_all(self, _doc_orig, _doc):

        """returns 0 in case of storaging"""

        return 0

def main():

    """runs everything"""

    text_proc = TextProcessor()
    generate_tasks = Generator(text_proc)
    saving = Storage(generate_tasks)
    saving.doc_orig.save('Path to file to save original text')
    saving.doc.save('Path to file to save tasks')

if __name__ == "__main__":
    main()
